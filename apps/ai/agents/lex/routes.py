import io
import re
import base64
import json
import uuid
from datetime import datetime
from typing import Literal

from fastapi import APIRouter
from pydantic import BaseModel, ConfigDict, Field

from core.llm import LLMClient
from core.rag import RAGService
from core.models import ChatRequest, ChatSyncResponse
from core.config import settings
from core.utils import strip_json_fences, safe_json_loads
from agents.lex.agent import LexAgent

router = APIRouter(prefix="/ai/lex", tags=["Lex"])

from agents.registry import register_agent

_llm = LLMClient()
_rag = RAGService()
_agent = LexAgent(_llm, _rag)
register_agent(_agent)


# ── Models ───────────────────────────────────────────────────────────────────

class IngestDocumentRequest(BaseModel):
    user_id: str
    organization_id: str = ""
    document_name: str
    document_type: str = "nda"
    document_url: str | None = None
    pdf_base64: str | None = None
    metadata: dict = Field(default_factory=dict)

    model_config = ConfigDict(
        json_schema_extra={
            "example": {
                "user_id": "user_123",
                "document_name": "Acme Corp NDA 2025",
                "document_type": "nda",
                "document_url": "https://pub-xxxx.r2.dev/documents/acme-nda-2025.pdf",
            }
        }
    )


class IngestDocumentResponse(BaseModel):
    source_id: str
    chunks_created: int
    page_count: int
    summary: str
    key_topics: list[str]
    document_type_detected: str
    tokens_used: int = 0
    model_used: str = ""


class AnalyzeContractRequest(BaseModel):
    user_id: str
    organization_id: str = ""
    source_id: str
    metadata: dict = Field(default_factory=dict)

    model_config = ConfigDict(
        json_schema_extra={
            "example": {
                "user_id": "user_123",
                "source_id": "doc_abc123",
            }
        }
    )


class ClauseRisk(BaseModel):
    clause: str
    risk: str
    severity: str        # low / medium / high / critical
    recommendation: str
    confidence: str | None = None   # high / medium / low
    basis: str | None = None        # one-sentence legal justification


class ScoreBreakdown(BaseModel):
    critical: int = 0
    high: int = 0
    medium: int = 0
    low: int = 0


class ObligationItem(BaseModel):
    action: str
    deadline: str | None = None
    condition: str | None = None
    consequence: str | None = None


class PartyObligations(BaseModel):
    party: str
    items: list[ObligationItem]


class AmbiguousClause(BaseModel):
    clause: str
    section: str | None = None
    issue: str
    interpretation: str


class NegotiationPoint(BaseModel):
    priority: str        # high / medium / low
    clause: str
    issue: str
    suggested_change: str


class ContractAnalysis(BaseModel):
    # Overview
    document_type: str
    parties: list[str]
    effective_date: str
    governing_law: str
    jurisdiction: str

    # Summary
    executive_summary: str
    risk_level: str      # low / medium / high / critical
    risk_score: int      # 1–10

    # Risk breakdown
    risks: list[ClauseRisk]
    unusual_clauses: list[str]
    missing_protections: list[str]

    # Clause-by-clause
    clause_breakdown: list[dict]  # {section, title, summary, risk_level, notes}

    # Key terms
    key_terms: dict

    # Obligations per party
    obligations: dict    # {"Party A name": [...], "Party B name": [...]}

    # Negotiation guidance
    negotiation_points: list[NegotiationPoint]

    # Verdict
    overall_assessment: str
    recommended_action: str  # sign / negotiate / reject / legal_review_required

    # Enhanced fields (optional — absent on older cached results)
    score_breakdown: ScoreBreakdown | None = None
    obligations_structured: list[PartyObligations] | None = None
    ambiguous_clauses: list[AmbiguousClause] | None = None


class AnalyzeContractResponse(BaseModel):
    analysis: ContractAnalysis
    tokens_used: int = 0
    model_used: str = ""


class QueryDocumentRequest(BaseModel):
    user_id: str
    organization_id: str = ""
    source_id: str
    query: str
    top_k: int = 5
    metadata: dict = Field(default_factory=dict)

    model_config = ConfigDict(
        json_schema_extra={
            "example": {
                "user_id": "user_123",
                "source_id": "doc_abc123",
                "query": "What are the termination conditions and notice periods?",
                "top_k": 5,
            }
        }
    )


class QueryDocumentChunk(BaseModel):
    content: str
    score: float
    metadata: dict = {}


class QueryDocumentResponse(BaseModel):
    answer: str
    sources: list[QueryDocumentChunk]
    tokens_used: int = 0
    model_used: str = ""


class DraftDocumentRequest(BaseModel):
    user_id: str
    organization_id: str = ""
    document_type: str
    requirements: str
    jurisdiction: str = "United States (Delaware)"
    additional_clauses: list[str] = []
    metadata: dict = Field(default_factory=dict)

    model_config = ConfigDict(
        json_schema_extra={
            "example": {
                "user_id": "user_123",
                "document_type": "mutual_nda",
                "requirements": "Mutual NDA between two SaaS companies for a potential partnership discussion. 2-year term, covers product roadmap and customer data.",
                "jurisdiction": "United States (Delaware)",
                "additional_clauses": ["data_protection", "ip_assignment_exclusion"],
            }
        }
    )


class DraftDocumentResponse(BaseModel):
    document: str
    review_notes: list[str]
    tokens_used: int = 0
    model_used: str = ""


class ExplainRequest(BaseModel):
    user_id: str
    organization_id: str = ""
    text: str
    context: str | None = None
    metadata: dict = Field(default_factory=dict)

    model_config = ConfigDict(
        json_schema_extra={
            "example": {
                "user_id": "user_123",
                "text": "The Receiving Party agrees to hold the Confidential Information in strict confidence and not to disclose it to any third party without the prior written consent of the Disclosing Party.",
                "context": "This is from an NDA we're about to sign with a potential investor",
            }
        }
    )


class ExplainResponse(BaseModel):
    explanation: str
    key_terms: dict
    related_concepts: list[str]
    practical_implications: list[str]
    tokens_used: int = 0
    model_used: str = ""


class LegalResearchRequest(BaseModel):
    user_id: str
    organization_id: str = ""
    query: str
    jurisdiction: str = "United States"
    legal_areas: list[str] = []
    metadata: dict = Field(default_factory=dict)

    model_config = ConfigDict(
        json_schema_extra={
            "example": {
                "user_id": "user_123",
                "query": "What are the GDPR requirements for obtaining valid consent from users in the EU?",
                "jurisdiction": "EU",
                "legal_areas": ["data_privacy", "consent"],
            }
        }
    )


class LegalResearchResponse(BaseModel):
    summary: str
    applicable_laws: list[str]
    key_requirements: list[str]
    relevant_cases: list[str]
    practical_guidance: list[str]
    jurisdiction_notes: str
    confidence_level: str
    tokens_used: int = 0
    model_used: str = ""


class ComplianceCheckRequest(BaseModel):
    user_id: str
    organization_id: str = ""
    description: str
    frameworks: list[str]
    business_context: str = ""
    metadata: dict = Field(default_factory=dict)

    model_config = ConfigDict(
        json_schema_extra={
            "example": {
                "user_id": "user_123",
                "description": "We store EU user email addresses and behavioral analytics data on AWS US-East servers with 90-day retention and no explicit consent flow.",
                "frameworks": ["GDPR", "CCPA"],
                "business_context": "B2B SaaS with EU and California customers",
            }
        }
    )


class ComplianceCheckResponse(BaseModel):
    overall_status: str
    framework_results: list[dict]
    critical_gaps: list[str]
    remediation_steps: list[dict]
    estimated_effort: str
    tokens_used: int = 0
    model_used: str = ""


class DeleteSourceRequest(BaseModel):
    user_id: str
    source_id: str


class DeleteSourceResponse(BaseModel):
    deleted_chunks: int


class ListSourcesRequest(BaseModel):
    user_id: str


class SourceSummary(BaseModel):
    source_id: str
    source_type: str
    source_agent: str
    metadata: dict = {}
    created_at: str | None = None


class ListSourcesResponse(BaseModel):
    sources: list[SourceSummary]


# ── Routes ───────────────────────────────────────────────────────────────────

@router.post("/chat", response_model=ChatSyncResponse, summary="Lex chat")
async def lex_chat(request: ChatRequest) -> ChatSyncResponse:
    """Get Lex's legal response."""
    return await _agent.chat_sync(request)


@router.post("/ingest-document", response_model=IngestDocumentResponse, summary="Ingest legal document")
async def ingest_document(request: IngestDocumentRequest) -> IngestDocumentResponse:
    """Upload and process a legal document (PDF) for RAG-powered analysis."""
    if settings.MOCK_MODE:
        source_id = f"doc_{str(uuid.uuid4())[:8]}"
        return IngestDocumentResponse(
            source_id=source_id,
            chunks_created=8,
            page_count=4,
            summary=(
                "This document is a Mutual Non-Disclosure Agreement between two parties. "
                "It covers confidentiality obligations, permitted disclosures, intellectual property ownership, "
                "and dispute resolution. Duration is 5 years with Delaware governing law."
            ),
            key_topics=["confidentiality", "intellectual_property", "non_solicitation", "term_and_termination", "dispute_resolution"],
            document_type_detected="mutual_nda",
        )

    import asyncio
    import base64 as _base64
    from fastapi import HTTPException

    if request.pdf_base64:
        try:
            pdf_bytes = _base64.b64decode(request.pdf_base64)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Invalid base64 PDF data: {e}")
    elif request.document_url:
        import httpx
        try:
            async with httpx.AsyncClient(timeout=settings.R2_FETCH_TIMEOUT) as client:
                resp = await client.get(request.document_url)
                resp.raise_for_status()
                pdf_bytes = resp.content
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to fetch document from URL: {e}")
    else:
        raise HTTPException(status_code=400, detail="Either pdf_base64 or document_url is required")

    from core.pdf_reader import extract_text_with_vision, extract_pages

    # Extract text once — reused for both ingestion and metadata
    try:
        full_text = await extract_text_with_vision(pdf_bytes, _llm)
    except Exception:
        from core.pdf_reader import extract_text
        full_text = extract_text(pdf_bytes)

    page_count = len(extract_pages(pdf_bytes))
    source_id = f"doc_{str(uuid.uuid4())[:8]}"
    meta = {"document_name": request.document_name, "document_type": request.document_type}

    async def _extract_metadata():
        raw = await _llm.complete(
            provider=_agent.default_provider,
            model=_agent.default_model,
            system="You are a legal document analyst. Be precise and concise.",
            messages=[{"role": "user", "content": (
                f"Analyze this legal document and return ONLY a JSON object (no markdown fences) with exactly these keys:\n"
                f"- summary: 2-3 sentence overview of what the document is, who the parties are, and its main purpose\n"
                f"- key_topics: list of 4-8 specific legal topic strings (e.g. confidentiality, ip_assignment, termination, governing_law)\n"
                f"- document_type_detected: one of nda, mou, employment_agreement, service_agreement, partnership_agreement, "
                f"shareholder_agreement, lease_agreement, loan_agreement, settlement_agreement, other\n\n"
                f"Document:\n{full_text[:12000]}"
            )}],
            max_tokens=512,
        )
        return raw

    # Run ingestion and metadata extraction in parallel
    chunks_count, meta_raw = await asyncio.gather(
        _rag.ingest(request.user_id, full_text, "pdf", source_id, "lex", meta),
        _extract_metadata(),
    )

    tokens_used = _llm.count_tokens(meta_raw)
    try:
        meta_data = json.loads(strip_json_fences(meta_raw))
        summary = meta_data.get("summary", "")
        key_topics = meta_data.get("key_topics", [])
        document_type_detected = meta_data.get("document_type_detected", request.document_type)
    except Exception:
        summary = meta_raw[:300]
        key_topics = []
        document_type_detected = request.document_type

    return IngestDocumentResponse(
        source_id=source_id,
        chunks_created=chunks_count,
        page_count=page_count,
        summary=summary,
        key_topics=key_topics,
        document_type_detected=document_type_detected,
        tokens_used=tokens_used,
        model_used=_agent.default_model,
    )


@router.post("/analyze-contract", response_model=AnalyzeContractResponse, summary="Analyze contract")
async def analyze_contract(request: AnalyzeContractRequest) -> AnalyzeContractResponse:
    """Fetch all chunks for source_id and perform a full structured contract analysis."""
    if settings.MOCK_MODE:
        return AnalyzeContractResponse(
            analysis=ContractAnalysis(
                document_type="Mutual Non-Disclosure Agreement",
                parties=["Acme Corp (Disclosing Party)", "Beta Inc (Receiving Party)"],
                effective_date="January 1, 2025",
                governing_law="Delaware General Corporation Law",
                jurisdiction="United States (Delaware)",
                executive_summary=(
                    "A mutual NDA between two technology companies for partnership evaluation. "
                    "The agreement is broadly written with several one-sided provisions favoring the disclosing party, "
                    "most notably a residuals clause that creates IP leakage risk."
                ),
                risk_level="medium",
                risk_score=6,
                risks=[
                    ClauseRisk(
                        clause="Definition of Confidential Information (Section 2)",
                        risk="Overly broad — includes all verbal communications with no follow-up written confirmation requirement, making scope unmanageable.",
                        severity="medium",
                        recommendation="Add a requirement that verbal disclosures be confirmed in writing within 10 business days to be considered Confidential Information.",
                        confidence="high",
                        basis="Courts require written confirmation of verbal disclosures to be enforceable under the Uniform Trade Secrets Act § 1(4)(i).",
                    ),
                    ClauseRisk(
                        clause="Residuals Clause (Section 7)",
                        risk="Allows the receiving party to use residual knowledge retained in unaided memory — creates IP leakage risk for your product roadmap and technical architecture.",
                        severity="high",
                        recommendation="Remove entirely or limit to generic industry knowledge, explicitly excluding product roadmap, source code, and customer data.",
                        confidence="high",
                        basis="Delaware courts have repeatedly invalidated residuals clauses lacking explicit scope limitations as incompatible with DTSA protections.",
                    ),
                    ClauseRisk(
                        clause="Duration (Section 5)",
                        risk="5-year term is 2× the industry standard of 2–3 years for startup-stage NDAs.",
                        severity="low",
                        recommendation="Renegotiate to 2–3 years with a survival clause limited to 1 year post-termination for specific categories.",
                        confidence="medium",
                        basis="Typical NDA enforceability period in Delaware startup practice is 2–3 years per consistent market precedent.",
                    ),
                    ClauseRisk(
                        clause="Unilateral Termination (Section 9)",
                        risk="Disclosing party can terminate with 30 days notice, but confidentiality obligations survive indefinitely — your obligations continue even after termination.",
                        severity="medium",
                        recommendation="Cap survival of obligations to 2 years post-termination and require mutual consent for early termination.",
                        confidence="high",
                        basis="Indefinite post-termination survival clauses have been found commercially unreasonable in multiple Delaware chancery court rulings.",
                    ),
                ],
                unusual_clauses=[
                    "Residuals clause (Section 7) — rare in mutual NDAs, typically found only in one-sided agreements favoring large enterprises",
                    "No limitation on injunctive relief scope — allows either party to seek unlimited injunctive relief without bond",
                    "No carve-out for information independently developed — standard protection is missing",
                ],
                missing_protections=[
                    "Dispute resolution mechanism — arbitration vs. litigation not specified",
                    "Data protection / GDPR compliance obligations — no mention despite potential EU data exchange",
                    "Limitation of liability — no cap on damages beyond injunctive relief",
                    "Carve-out for information independently developed without reference to confidential information",
                    "Return or destruction of materials clause upon termination",
                ],
                clause_breakdown=[
                    {"section": "1", "title": "Purpose", "summary": "Defines the scope of the relationship as partnership evaluation.", "risk_level": "low", "notes": "Standard. No issues."},
                    {"section": "2", "title": "Definition of Confidential Information", "summary": "Broadly defines all information, including verbal disclosures, as confidential.", "risk_level": "medium", "notes": "Verbal inclusion without written confirmation requirement is problematic."},
                    {"section": "3", "title": "Obligations", "summary": "Standard confidentiality obligations — hold in confidence, restrict access, use only for Purpose.", "risk_level": "low", "notes": "Standard. No issues."},
                    {"section": "5", "title": "Term", "summary": "5-year confidentiality period from Effective Date.", "risk_level": "low", "notes": "Above industry standard. Negotiate down."},
                    {"section": "7", "title": "Residuals", "summary": "Permits use of retained knowledge in unaided memory for any purpose.", "risk_level": "high", "notes": "Significant IP risk. Should be removed or heavily restricted."},
                    {"section": "9", "title": "Termination", "summary": "Either party may terminate with 30 days written notice. Obligations survive.", "risk_level": "medium", "notes": "Indefinite survival of obligations post-termination is unusual."},
                ],
                key_terms={
                    "duration": "5 years",
                    "governing_law": "Delaware, United States",
                    "scope": "Product roadmap, customer data, financial projections, technical architecture",
                    "permitted_disclosures": "Legal counsel and advisors who are themselves bound by NDA",
                    "dispute_resolution": "Not specified",
                    "termination_notice": "30 days written notice",
                    "survival": "Indefinite post-termination",
                },
                obligations={
                    "Acme Corp": [
                        "Hold Beta's Confidential Information in strict confidence",
                        "Limit access to employees with need-to-know",
                        "Use information solely for the Purpose",
                        "Notify Beta immediately of any unauthorized disclosure",
                    ],
                    "Beta Inc": [
                        "Hold Acme's Confidential Information in strict confidence",
                        "Limit access to employees with need-to-know",
                        "Use information solely for the Purpose",
                        "Notify Acme immediately of any unauthorized disclosure",
                    ],
                },
                negotiation_points=[
                    NegotiationPoint(
                        priority="high",
                        clause="Residuals Clause (Section 7)",
                        issue="Permits unrestricted use of retained knowledge — directly undermines the NDA's purpose.",
                        suggested_change="Delete Section 7 entirely, or restrict to generic industry knowledge explicitly excluding product IP, source code, and customer data.",
                    ),
                    NegotiationPoint(
                        priority="medium",
                        clause="Definition of Confidential Information (Section 2)",
                        issue="Verbal communications without written confirmation create enforcement ambiguity.",
                        suggested_change="Add: 'Verbal disclosures must be confirmed in writing within 10 business days to constitute Confidential Information.'",
                    ),
                    NegotiationPoint(
                        priority="medium",
                        clause="Termination / Survival (Section 9)",
                        issue="Indefinite survival of obligations is commercially unusual and burdensome.",
                        suggested_change="Cap survival at 3 years post-termination for general confidentiality; allow indefinite survival only for trade secrets.",
                    ),
                    NegotiationPoint(
                        priority="low",
                        clause="Duration (Section 5)",
                        issue="5-year term exceeds market standard.",
                        suggested_change="Reduce to 2–3 years.",
                    ),
                ],
                overall_assessment=(
                    "This NDA is weighted toward the disclosing party and contains one high-risk clause (residuals) "
                    "that should be a hard blocker. The agreement is otherwise workable with targeted modifications. "
                    "Do not sign as-is."
                ),
                recommended_action="negotiate",
                score_breakdown=ScoreBreakdown(critical=0, high=1, medium=2, low=1),
                obligations_structured=[
                    PartyObligations(
                        party="Acme Corp",
                        items=[
                            ObligationItem(action="Hold Beta's Confidential Information in strict confidence", deadline=None, condition=None, consequence="Injunctive relief and damages"),
                            ObligationItem(action="Limit access to employees with need-to-know", deadline=None, condition=None, consequence=None),
                            ObligationItem(action="Use information solely for the Purpose", deadline=None, condition=None, consequence="Termination of agreement"),
                            ObligationItem(action="Notify Beta immediately of any unauthorized disclosure", deadline="Immediately upon discovery", condition="Unauthorized disclosure occurs", consequence="Breach of agreement if not notified"),
                        ],
                    ),
                    PartyObligations(
                        party="Beta Inc",
                        items=[
                            ObligationItem(action="Hold Acme's Confidential Information in strict confidence", deadline=None, condition=None, consequence="Injunctive relief and damages"),
                            ObligationItem(action="Return or destroy all Confidential Information upon request", deadline="10 business days after written request", condition="Either party terminates or requests return", consequence="Material breach if not complied"),
                        ],
                    ),
                ],
                ambiguous_clauses=[
                    AmbiguousClause(
                        clause="reasonable efforts",
                        section="Section 3",
                        issue="'Reasonable efforts' is legally undefined and courts apply it inconsistently — some treat it as equivalent to 'best efforts' requiring maximum exertion; others treat it as a lower standard.",
                        interpretation="Delaware courts generally treat 'reasonable efforts' as less demanding than 'best efforts' but require demonstrable good-faith steps toward the obligation's objective.",
                    ),
                    AmbiguousClause(
                        clause="promptly",
                        section="Section 8",
                        issue="'Promptly' has no defined timeframe in this agreement, creating ambiguity for notice and cure obligations.",
                        interpretation="Without a defined period, Delaware courts apply a reasonableness standard — typically interpreted as within 3–5 business days for notice obligations in commercial contracts.",
                    ),
                ],
            ),
        )

    chunks = await _rag.retrieve_by_source(request.user_id, request.source_id)
    if not chunks:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail=f"No document found for source_id '{request.source_id}'")

    full_text = "\n\n".join(c.get("content", "") for c in chunks)

    system = await _agent.build_system_prompt(request.user_id, request.organization_id, use_brand_kit=False)
    memory_context = request.metadata.get("memory_context", "")
    if memory_context:
        system += f"\n\n## Memory Context\n{memory_context}"
    raw = await _llm.complete(
        provider=_agent.default_provider,
        model=_agent.default_model,
        system=system,
        messages=[{"role": "user", "content": (
            f"Perform a thorough, senior-attorney-level legal analysis of this contract. "
            f"Be specific — quote exact clause language when relevant, name exact section numbers, "
            f"and explain practical real-world impact, not just legal theory.\n\n"
            f"CONTRACT:\n{full_text}\n\n"
            "Return ONLY a valid JSON object (no markdown fences, no commentary) with EXACTLY these keys:\n\n"
            "document_type (string — precise document type, e.g. 'Mutual Non-Disclosure Agreement'),\n"
            "parties (list of strings — each entry: 'Full Legal Name (Role)', e.g. 'Acme Corp (Disclosing Party)'),\n"
            "effective_date (string — exact date or 'Not specified'),\n"
            "governing_law (string — state/country law that governs),\n"
            "jurisdiction (string — courts where disputes must be filed),\n"
            "executive_summary (string — 4–6 sentences: what the document is, who the parties are, "
            "its main commercial purpose, overall balance/fairness, and your plain-English verdict on "
            "whether it is founder-friendly or lopsided),\n"
            "risk_level (string — one of: low/medium/high/critical),\n"
            "risk_score (integer 1–10 where 1=essentially no risk, 10=do not sign),\n"
            "score_breakdown (object — count of risks at each severity level, fields: critical (int), high (int), medium (int), low (int) — must sum to the total number of risks),\n"
            "risks (list of objects — identify ALL material risks, minimum 3, with fields: "
            "clause (exact section name/number), "
            "risk (specific problem and its practical business impact — at least 2 sentences), "
            "severity (low/medium/high/critical), "
            "recommendation (specific actionable fix — proposed language change or deletion), "
            "confidence (high/medium/low — how certain you are this is a real enforceable risk based on case law or statute), "
            "basis (string — one sentence citing the specific legal authority or market precedent that supports this risk rating)),\n"
            "unusual_clauses (list of strings — clauses that deviate from market standard; for each, "
            "name the section and explain what is unusual and why it matters),\n"
            "missing_protections (list of strings — standard protections absent from this agreement; "
            "for each, name what is missing and the risk that creates),\n"
            "clause_breakdown (list of objects — analyze EVERY numbered section/article; fields: "
            "section (number/letter), title (clause title), "
            "summary (2–3 sentences on what it does and its effect), "
            "risk_level (low/medium/high/critical), "
            "notes (specific issues, unusual language, or 'Standard — no issues')),\n"
            "key_terms (dict string->string — important defined terms and their practical meaning, "
            "minimum 5 entries),\n"
            "obligations (dict — party_name -> list of specific obligation strings; be exhaustive, "
            "list every obligation each party takes on),\n"
            "obligations_structured (list of objects — same obligations in structured form, each: "
            "party (string — exact party name matching the parties list), "
            "items (list of objects, each with: "
            "action (string — what the party must do), "
            "deadline (string or null — human-readable deadline e.g. '30 days after termination', null if none), "
            "condition (string or null — condition that triggers this obligation, null if unconditional), "
            "consequence (string or null — consequence of non-performance, null if unspecified))),\n"
            "ambiguous_clauses (list of objects — flag any legally vague or undefined terms that could create disputes; "
            "focus on phrases like 'reasonable efforts', 'material adverse change', 'promptly', 'good faith', "
            "'commercially reasonable', or any capitalized Defined Term that is used but not defined in this agreement; "
            "fields: clause (string — the exact vague phrase or undefined term), "
            "section (string or null — section number/name where it appears, null if it appears in multiple places), "
            "issue (string — why this language creates legal uncertainty), "
            "interpretation (string — how courts in the governing jurisdiction typically interpret this language)),\n"
            "negotiation_points (list of objects — prioritized redline targets, minimum 3, fields: "
            "priority (high/medium/low), clause (section name/number), "
            "issue (what is wrong and why it must change), "
            "suggested_change (exact proposed language or specific deletion instruction)),\n"
            "overall_assessment (string — 3–4 sentences: concrete verdict on who this agreement "
            "favors, what would change your recommendation, and whether to sign as-is),\n"
            "recommended_action (string — one of: sign/negotiate/reject/legal_review_required)"
        )}],
        max_tokens=10000,
    )
    tokens_used = _llm.count_tokens(raw)
    try:
        data = json.loads(strip_json_fences(raw))
        risks = [ClauseRisk(**r) for r in data.pop("risks", [])]
        negotiation_points = [NegotiationPoint(**n) for n in data.pop("negotiation_points", [])]
        raw_obligations_structured = data.pop("obligations_structured", [])
        obligations_structured = [
            PartyObligations(party=p["party"], items=[ObligationItem(**item) for item in p.get("items", [])])
            for p in raw_obligations_structured
        ] if raw_obligations_structured else None
        raw_ambiguous = data.pop("ambiguous_clauses", [])
        ambiguous_clauses = [AmbiguousClause(**c) for c in raw_ambiguous] if raw_ambiguous else None
        raw_breakdown = data.pop("score_breakdown", None)
        score_breakdown = ScoreBreakdown(**raw_breakdown) if raw_breakdown else None
        analysis = ContractAnalysis(
            **data,
            risks=risks,
            negotiation_points=negotiation_points,
            obligations_structured=obligations_structured,
            ambiguous_clauses=ambiguous_clauses,
            score_breakdown=score_breakdown,
        )
    except Exception:
        analysis = ContractAnalysis(
            document_type="Unknown",
            parties=[], effective_date="", governing_law="", jurisdiction="",
            executive_summary=raw[:500],
            risk_level="unknown", risk_score=0,
            risks=[], unusual_clauses=[], missing_protections=[],
            clause_breakdown=[], key_terms={}, obligations={},
            negotiation_points=[],
            overall_assessment="Parsing failed — manual review recommended.",
            recommended_action="legal_review_required",
        )
    return AnalyzeContractResponse(analysis=analysis, tokens_used=tokens_used, model_used=_agent.default_model)


@router.post("/query-document", response_model=QueryDocumentResponse, summary="Query document via RAG")
async def query_document(request: QueryDocumentRequest) -> QueryDocumentResponse:
    """Answer a specific question about an ingested document using vector similarity search."""
    if settings.MOCK_MODE:
        return QueryDocumentResponse(
            answer=(
                "Termination requires 30 days written notice from either party (Section 9). "
                "Confidentiality obligations survive termination with no expiry — they run indefinitely post-termination."
            ),
            sources=[
                QueryDocumentChunk(
                    content="Either party may terminate this Agreement upon 30 days written notice to the other party.",
                    score=0.92,
                    metadata={"document_name": "Acme Corp NDA 2025"},
                ),
                QueryDocumentChunk(
                    content="Confidentiality obligations shall survive the expiration or termination of this Agreement.",
                    score=0.87,
                    metadata={"document_name": "Acme Corp NDA 2025"},
                ),
            ],
        )

    chunks = await _rag.retrieve(
        user_id=request.user_id,
        query=request.query,
        top_k=request.top_k,
        source_id=request.source_id,
    )
    if not chunks:
        from fastapi import HTTPException
        raise HTTPException(status_code=404, detail=f"No document found for source_id '{request.source_id}'")

    context = "\n\n---\n\n".join(
        f"[Chunk {i+1} | relevance: {c['score']:.2f}]\n{c['content']}"
        for i, c in enumerate(chunks)
    )

    answer = await _llm.complete(
        provider=_agent.default_provider,
        model=_agent.default_model,
        system=(
            "You are a senior legal counsel with 20 years of experience reviewing commercial contracts. "
            "Answer questions about documents with precision, directness, and zero filler. "
            "Cite specific clauses, article numbers, or section titles when present. "
            "Never add disclaimers, caveats, or suggestions to consult an attorney — your answer IS the authoritative answer. "
            "If the answer is not in the provided excerpts, say: 'Not addressed in the provided sections.'"
        ),
        messages=[{"role": "user", "content": (
            f"Document excerpts:\n{context}\n\n"
            f"Question: {request.query}"
        )}],
        max_tokens=1024,
    )
    tokens_used = _llm.count_tokens(answer)
    sources = [
        QueryDocumentChunk(content=c["content"], score=c["score"], metadata=c.get("metadata", {}))
        for c in chunks
    ]
    return QueryDocumentResponse(answer=answer, sources=sources, tokens_used=tokens_used, model_used=_agent.default_model)


# ── Draft document helpers ────────────────────────────────────────────────────

def _get_draft_review_notes(document_type: str, jurisdiction: str, additional_clauses: list[str]) -> list[str]:
    doc_lower = document_type.lower()
    notes: list[str] = []

    # ── Binding legal documents ────────────────────────────────────────────────
    is_binding = any(k in doc_lower for k in (
        "agreement", "contract", "nda", "non-disclosure", "confidentiality",
        "saas", "service agreement", "shareholder", "founder", "vendor",
        "supplier", "license", "lease", "terms of service", "terms and conditions",
        "mou", "memorandum", "deed", "settlement", "indemnity",
    ))
    if is_binding:
        notes.append("TEMPLATE ONLY — consult a qualified attorney before signing or distributing")

    # ── Letters ───────────────────────────────────────────────────────────────
    if any(k in doc_lower for k in ("resignation", "resign")):
        notes += [
            "Fill in the employee's full last name and any remaining [BRACKET] fields",
            "Confirm the exact last working day — check your employment contract's required notice period",
            "Review the reason stated — consider whether to keep, soften, or omit specific reasons before sending",
            "Verify whether HR requires a separate clearance or handover checklist",
            "Keep a signed or sent copy for your personal records",
        ]
    elif any(k in doc_lower for k in ("offer letter", "job offer")):
        notes += [
            "Confirm all compensation figures (salary, bonus) are accurate before sending",
            "Verify the start date, role title, and reporting line",
            "Reference your equity/option plan document if applicable",
            "Have HR or legal review before sending to the candidate",
        ]
    elif any(k in doc_lower for k in ("demand letter", "legal notice", "cease and desist")):
        notes += [
            "Verify the recipient's full legal name and mailing address",
            "Confirm all factual claims are accurate and can be documented",
            "Set a realistic response deadline (typically 10–30 days)",
            "Send via certified mail and retain proof of delivery",
            "Have an attorney review before sending — a poorly worded demand can undermine your position",
        ]
    elif "cover letter" in doc_lower:
        notes += [
            "Personalise the opening paragraph to reference the specific role and company",
            "Ensure your contact details and the hiring manager's name are correct",
            "Tailor the skills highlighted to match the job description",
        ]
    elif any(k in doc_lower for k in ("recommendation letter", "reference letter")):
        notes += [
            "Confirm the subject's name, role, and dates of association are correct",
            "Add specific examples or achievements to make the letter more credible",
            "Include your title and contact details so the recipient can verify",
        ]

    # ── Contracts & agreements ────────────────────────────────────────────────
    elif any(k in doc_lower for k in ("nda", "non-disclosure", "confidentiality")):
        notes += [
            "Fill in full legal names and entity types (LLC, Inc., Ltd.) for both parties",
            "Specify the exact confidentiality period — standard range is 2–5 years",
            "Define the scope of 'Confidential Information' to match what you'll actually share",
            "Add a DPA clause if sharing personal data (GDPR/CCPA may apply)",
            "Confirm the governing law matches where your business is registered",
        ]
    elif any(k in doc_lower for k in ("employment contract", "employment agreement")):
        notes += [
            "Confirm compensation, start date, and role title are accurate",
            "Verify IP assignment clause compliance with your state's law",
            "Review non-compete scope — many US states (CA, MN, ND) restrict or ban them",
            "Confirm at-will / probation language matches your local employment law",
            "Add equity or bonus details referencing a separate option plan if applicable",
        ]
    elif any(k in doc_lower for k in ("saas", "software", "service agreement", "subscription")):
        notes += [
            "Confirm SLA uptime percentage and credit/remedy structure",
            "Check liability cap — typically tied to fees paid in the prior 12 months",
            "Add a DPA if you handle EU/UK personal data",
            "Verify auto-renewal window and the required cancellation notice period",
        ]
    elif any(k in doc_lower for k in ("founder", "co-founder", "shareholder")):
        notes += [
            "Confirm equity split percentages and vesting schedule (common: 4-year with 1-year cliff)",
            "Review IP assignment — all pre-existing relevant IP must be assigned to the company",
            "Verify drag-along and tag-along rights align with your cap table intentions",
            "Each party should have this reviewed by independent counsel before signing",
        ]
    elif any(k in doc_lower for k in ("vendor", "supplier", "procurement", "purchase order")):
        notes += [
            "Confirm payment terms (Net 30/60) and any late payment penalty rate",
            "Review indemnification — ensure it is mutual and proportionate",
            "Add acceptance testing criteria for software or physical deliverables",
            "Confirm warranty period and remedy for defective deliverables",
        ]

    # ── Notices, memos, policies ──────────────────────────────────────────────
    elif any(k in doc_lower for k in ("notice", "memo", "memorandum", "circular")):
        notes += [
            "Verify the To / From / Date fields are filled in correctly",
            "Confirm the subject line accurately describes the purpose",
            "Review the action items or deadlines stated and ensure they are achievable",
        ]
    elif any(k in doc_lower for k in ("privacy policy", "terms of service", "cookie policy", "refund policy")):
        notes += [
            "Have a lawyer review before publishing — policy documents can create legal obligations",
            "Confirm data types collected and retention periods are accurately described",
            "Verify compliance with applicable regulations (GDPR, CCPA, etc.) for your user base",
            "Keep a version history and update the 'last updated' date whenever you change the policy",
        ]

    # ── Generic fallback ──────────────────────────────────────────────────────
    else:
        notes += [
            "Fill in all [BRACKET] placeholders before using or sending this document",
            "Review all names, dates, and factual details for accuracy",
            "Adjust the tone or any specific language to fit your exact situation",
        ]
        if is_binding:
            notes += [
                "Verify the governing law and jurisdiction clause match your intended forum",
                "Review any liability, indemnification, or penalty clauses carefully",
            ]

    if additional_clauses:
        notes.append(
            f"Requested elements ({', '.join(additional_clauses[:3])}) have been included — verify they fit your situation"
        )

    if is_binding:
        if "india" in jurisdiction.lower():
            notes.append("Indian law document — confirm stamp duty requirements for your state before execution")
        elif "uk" in jurisdiction.lower() or "england" in jurisdiction.lower():
            notes.append("UK law document — confirm post-Brexit cross-border implications if any EU parties are involved")
        elif "eu" in jurisdiction.lower() or "europe" in jurisdiction.lower():
            notes.append("EU-governed document — GDPR Article 28 DPA is likely required; verify with your DPO")

    return notes


def _parse_document_sections(text: str) -> list[dict]:
    """Classify each line of a legal document for structured rendering."""
    lines = text.split("\n")
    sections: list[dict] = []
    in_signature = False
    found_title = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            sections.append({"type": "blank", "text": ""})
            continue

        if re.match(r"^(IN WITNESS WHEREOF|SIGNED BY|SIGNATURE PAGE|EXECUTED AS OF)", stripped, re.IGNORECASE):
            in_signature = True

        if in_signature:
            sections.append({"type": "signature", "text": stripped})
            continue

        if re.match(r"^\d+(\.\d+)*\.?\s+[A-Z]", stripped):
            sections.append({"type": "heading", "text": stripped})
            continue

        is_upper = stripped == stripped.upper() and re.search(r"[A-Z]", stripped) and len(stripped) > 3
        if is_upper and not found_title:
            sections.append({"type": "title", "text": stripped})
            found_title = True
            continue
        if is_upper and len(stripped.split()) <= 10:
            sections.append({"type": "heading", "text": stripped})
            continue

        if re.match(r"^\([a-z]\)\s", stripped):
            sections.append({"type": "subitem", "text": stripped})
            continue

        sections.append({"type": "body", "text": stripped})

    return sections


def _generate_docx(document_text: str, document_type: str) -> bytes:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument()
    for section in doc.sections:
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    parsed = _parse_document_sections(document_text)
    sig_lines: list[str] = []

    for s in parsed:
        if s["type"] == "blank":
            continue
        elif s["type"] == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(s["text"])
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
            p.paragraph_format.space_after = Pt(6)
        elif s["type"] == "heading":
            p = doc.add_paragraph()
            run = p.add_run(s["text"])
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif s["type"] == "subitem":
            p = doc.add_paragraph(s["text"])
            p.paragraph_format.left_indent = Inches(0.3)
            p.runs[0].font.name = "Times New Roman"
        elif s["type"] == "signature":
            sig_lines.append(s["text"])
        else:
            p = doc.add_paragraph(s["text"])
            p.paragraph_format.space_after = Pt(4)
            if p.runs:
                p.runs[0].font.name = "Times New Roman"

    if sig_lines:
        doc.add_paragraph()
        for line in sig_lines:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            if p.runs:
                p.runs[0].font.name = "Times New Roman"

    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.text = ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_pdf(document_text: str, document_type: str) -> bytes:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

    def _esc(t: str) -> str:
        return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    buf = io.BytesIO()
    doc_template = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=1.25 * inch, rightMargin=1.25 * inch,
        topMargin=1.0 * inch, bottomMargin=0.9 * inch,
    )

    base = getSampleStyleSheet()["Normal"]
    title_style = ParagraphStyle("LexTitle", parent=base, fontName="Times-Bold",
                                  fontSize=14, alignment=TA_CENTER, spaceAfter=10)
    heading_style = ParagraphStyle("LexHeading", parent=base, fontName="Times-Bold",
                                    fontSize=12, alignment=TA_LEFT, spaceBefore=12, spaceAfter=5)
    body_style = ParagraphStyle("LexBody", parent=base, fontName="Times-Roman",
                                 fontSize=12, alignment=TA_JUSTIFY, spaceAfter=5, leading=17)
    sub_style = ParagraphStyle("LexSub", parent=base, fontName="Times-Roman",
                                fontSize=12, alignment=TA_LEFT, leftIndent=18, spaceAfter=4, leading=17)

    parsed = _parse_document_sections(document_text)
    story = []
    sig_lines: list[str] = []

    for s in parsed:
        if s["type"] == "blank":
            story.append(Spacer(1, 5))
        elif s["type"] == "title":
            story.append(Paragraph(_esc(s["text"]), title_style))
        elif s["type"] == "heading":
            story.append(Paragraph(_esc(s["text"]), heading_style))
        elif s["type"] == "subitem":
            story.append(Paragraph(_esc(s["text"]), sub_style))
        elif s["type"] == "signature":
            sig_lines.append(s["text"])
        else:
            story.append(Paragraph(_esc(s["text"]), body_style))

    if sig_lines:
        story.append(Spacer(1, 20))
        for line in sig_lines:
            story.append(Paragraph(_esc(line), body_style))

    def _footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Times-Italic", 8)
        canvas.drawRightString(LETTER[0] - 1.25 * inch, 0.5 * inch, f"Page {doc.page}")
        canvas.restoreState()

    doc_template.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buf.getvalue()


@router.post("/draft-document", response_model=DraftDocumentResponse, summary="Draft legal document")
async def draft_document(request: DraftDocumentRequest) -> DraftDocumentResponse:
    """Draft a legal document template based on requirements."""
    if settings.MOCK_MODE:
        return DraftDocumentResponse(
            document=f"""MUTUAL NON-DISCLOSURE AGREEMENT

This Mutual Non-Disclosure Agreement ("Agreement") is entered into as of [DATE] by and between:

Party A: [COMPANY NAME], a Delaware corporation ("Company A")
Party B: [COMPANY NAME], a Delaware corporation ("Company B")

(Each a "Party" and collectively the "Parties")

1. PURPOSE
The Parties wish to explore a potential business relationship ("Purpose") and may need to disclose certain confidential information to each other.

2. DEFINITION OF CONFIDENTIAL INFORMATION
"Confidential Information" means any information disclosed by either Party to the other Party, either directly or indirectly, in writing, orally, or by inspection of tangible objects, that is designated as "Confidential" at the time of disclosure.

Confidential Information does NOT include information that:
(a) Is or becomes publicly available through no breach of this Agreement
(b) Was rightfully known to the Receiving Party before disclosure
(c) Is independently developed by the Receiving Party without use of Confidential Information
(d) Is required to be disclosed by law or court order

3. OBLIGATIONS
Each Party agrees to: (a) hold the other's Confidential Information in strict confidence; (b) not disclose it to third parties without prior written consent; (c) use it solely for the Purpose; (d) limit access to employees with a need to know.

4. TERM
This Agreement shall remain in effect for two (2) years from the Effective Date, unless earlier terminated by either Party with 30 days written notice.

5. GOVERNING LAW
This Agreement shall be governed by the laws of the State of {request.jurisdiction.split('(')[-1].rstrip(')') if '(' in request.jurisdiction else 'Delaware'}, without regard to its conflict of laws provisions.

6. ENTIRE AGREEMENT
This Agreement constitutes the entire agreement between the Parties with respect to the subject matter hereof.

IN WITNESS WHEREOF, the Parties have executed this Agreement as of the date first written above.

COMPANY A                               COMPANY B
By: _______________________             By: _______________________
Name: _____________________            Name: _____________________
Title: _____________________           Title: _____________________
Date: ______________________           Date: ______________________
""",
            review_notes=[
                "TEMPLATE ONLY – requires customization before use",
                "Insert full legal names and entity types in the header",
                "Add specific description of the business purpose in Section 1",
                "Adjust the definition scope in Section 2 based on your actual sharing needs",
                "Consider adding a data protection clause if sharing personal data (GDPR/CCPA implications)",
            ],
        )

    system = await _agent.build_system_prompt(request.user_id, request.organization_id)
    memory_context = request.metadata.get("memory_context", "")
    if memory_context:
        system += f"\n\n## Memory Context\n{memory_context}"
    doc_prompt = (
        f"Draft a complete, professional {request.document_type}.\n\n"
        f"REQUIREMENTS:\n{request.requirements}\n\n"
        f"JURISDICTION: {request.jurisdiction}\n"
        f"ADDITIONAL ELEMENTS: {', '.join(request.additional_clauses) if request.additional_clauses else 'None'}\n\n"
        "CRITICAL RULE: Use the format that is natural and standard for this specific document type.\n"
        "Do NOT impose a legal contract structure on every document. Match the format to what the document actually is:\n\n"
        "• Letter (resignation, offer, demand, cover, recommendation):\n"
        "  - Standard letter format: sender info, date, recipient info, subject line, greeting, body paragraphs, closing, signature\n"
        "  - Do NOT use numbered clauses. Do NOT write 'by and between' preambles. Write naturally like a real letter.\n\n"
        "• Legal agreement or contract (NDA, employment contract, SaaS agreement, shareholder agreement, vendor contract):\n"
        "  - Title in ALL CAPS, 'by and between' party preamble, numbered sections with ALL-CAPS headings\n"
        "  - Full two-party signature block at the end\n\n"
        "• Policy document (privacy policy, terms of service, code of conduct, refund policy):\n"
        "  - Numbered or headed sections, formal but readable prose, no bilateral party preamble\n\n"
        "• Notice or memo (termination notice, eviction notice, board resolution, HR memo):\n"
        "  - Appropriate header (To / From / Date / Re:), concise factual body, single-signature block if needed\n\n"
        "• Any other document: use whatever format professionals in that field actually use.\n\n"
        "ALWAYS:\n"
        "- Write complete, substantive content — not stubs or one-liners\n"
        "- Use [BRACKETS] only for specific details the user must fill in (names, dates, amounts)\n"
        "- Output ONLY the document itself — no intro line, no closing comment, no markdown fences"
    )
    raw = await _llm.complete(
        provider=_agent.default_provider, model=_agent.default_model,
        system=system,
        messages=[{"role": "user", "content": doc_prompt}],
        max_tokens=4096,
    )
    document_text = raw.strip()
    tokens_used = _llm.count_tokens(document_text)
    return DraftDocumentResponse(
        document=document_text,
        review_notes=_get_draft_review_notes(
            request.document_type, request.jurisdiction, request.additional_clauses
        ),
        tokens_used=tokens_used,
        model_used=_agent.default_model,
    )


# ── Export document ───────────────────────────────────────────────────────────

class ExportDocumentRequest(BaseModel):
    document: str
    format: Literal["docx", "pdf"]
    document_type: str = "Legal Document"


class ExportDocumentResponse(BaseModel):
    file_b64: str
    mime_type: str
    filename: str


@router.post("/export-document", response_model=ExportDocumentResponse, summary="Export document as DOCX or PDF")
async def export_document(request: ExportDocumentRequest) -> ExportDocumentResponse:
    """Convert a drafted legal document to a formatted DOCX or PDF binary."""
    safe_name = re.sub(r"[^\w\s-]", "", request.document_type).strip().replace(" ", "_").lower() or "legal_document"

    if request.format == "docx":
        binary = _generate_docx(request.document, request.document_type)
        return ExportDocumentResponse(
            file_b64=base64.b64encode(binary).decode(),
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=f"lex_{safe_name}.docx",
        )
    else:
        binary = _generate_pdf(request.document, request.document_type)
        return ExportDocumentResponse(
            file_b64=base64.b64encode(binary).decode(),
            mime_type="application/pdf",
            filename=f"lex_{safe_name}.pdf",
        )


@router.post("/explain", response_model=ExplainResponse, summary="Explain legal text")
async def explain_legal_text(request: ExplainRequest) -> ExplainResponse:
    """Explain legal text in plain English with practical implications."""
    if settings.MOCK_MODE:
        return ExplainResponse(
            explanation=(
                "In plain English: You must keep the other party's confidential information completely private "
                "and not share it with anyone else. You can only get their permission to share it, and that "
                "permission must be in writing. This is a standard confidentiality obligation found in most NDAs.\n\n"
                "What this means for you: Once you sign this, sharing anything they've told you (product plans, "
                "financials, customer lists, etc.) with anyone else – even your own investors or advisors – "
                "requires their explicit written approval first."
            ),
            key_terms={
                "Receiving Party": "The person or company receiving the confidential information (likely you)",
                "Confidential Information": "Any information marked or identified as confidential during disclosure",
                "Third Party": "Anyone other than the two parties signing this agreement",
                "Prior Written Consent": "Written permission obtained BEFORE the disclosure, not after",
                "Disclosing Party": "The party sharing their confidential information with you",
            },
            related_concepts=[
                "Trade secret protection",
                "Attorney-client privilege (exception to NDA obligations)",
                "Carve-outs for publicly available information",
                "Return or destruction of confidential information upon termination",
            ],
            practical_implications=[
                "You cannot share their information with your investors without getting written approval first",
                "Internal team members who don't need to know should NOT be briefed on their confidential details",
                "Keep records of all confidential disclosures – document what was shared and when",
                "If you're subpoenaed, you may have a legal obligation to disclose – notify the disclosing party immediately",
                "Verbal disclosures also count if marked confidential at the time",
            ],
        )

    system = await _agent.build_system_prompt(request.user_id, request.organization_id, use_brand_kit=False)
    memory_context = request.metadata.get("memory_context", "")
    if memory_context:
        system += f"\n\n## Memory Context\n{memory_context}"
    raw = await _llm.complete(
        provider=_agent.default_provider, model=_agent.default_model,
        system=system,
        messages=[{"role": "user", "content": (
            f"Explain this legal text in plain English:\n\n{request.text}\n\n"
            f"Context: {request.context or 'None'}\n\n"
            "Return JSON with fields: explanation, key_terms (dict), related_concepts (list), practical_implications (list)"
        )}],
    )
    tokens_used = _llm.count_tokens(raw)
    try:
        data = json.loads(strip_json_fences(raw))
        return ExplainResponse(**data, tokens_used=tokens_used, model_used=_agent.default_model)
    except Exception:
        return ExplainResponse(
            explanation=raw,
            key_terms={},
            related_concepts=[],
            practical_implications=[],
            tokens_used=tokens_used,
            model_used=_agent.default_model,
        )


@router.post("/legal-research", response_model=LegalResearchResponse, summary="Research legal questions")
async def legal_research(request: LegalResearchRequest) -> LegalResearchResponse:
    """Research laws, regulations, and case precedents for a legal question."""
    if settings.MOCK_MODE:
        return LegalResearchResponse(
            summary=(
                "GDPR Article 7 sets strict requirements for consent. Consent must be freely given, specific, "
                "informed, and unambiguous. Pre-ticked boxes, bundled consent, and vague language are explicitly prohibited. "
                "Organizations must be able to demonstrate that consent was properly obtained."
            ),
            applicable_laws=[
                "GDPR Article 7 (Conditions for consent)",
                "GDPR Recital 32 (Affirmative consent requirement)",
                "GDPR Article 4(11) (Definition of consent)",
                "ePrivacy Directive (for cookies and electronic communications)",
                "GDPR Article 17 (Right to erasure — consent withdrawal triggers this)",
            ],
            key_requirements=[
                "Consent must be a clear affirmative act (no pre-ticked boxes or silence)",
                "Granular consent required — separate consent per distinct processing purpose",
                "As easy to withdraw consent as to give it",
                "Maintain records of consent: timestamp, IP address, consent version, method",
                "Children under 16 require parental consent (age threshold varies by EU member state: 13-16)",
                "No bundled consent — consent for services cannot be conditioned on unrelated data processing",
            ],
            relevant_cases=[
                "Planet49 GmbH v Bundesverband (CJEU C-673/17, 2019) — pre-ticked boxes invalid",
                "Fashion ID GmbH v Verbraucherzentrale NRW (CJEU C-40/17, 2019) — joint controller liability",
                "CNIL enforcement actions (France, 2022-2023) — cookie consent dark patterns fined",
                "DSK (Germany) — guidance on valid consent for analytics tracking",
            ],
            practical_guidance=[
                "Implement a proper Consent Management Platform (CMP) with granular opt-in toggles",
                "Store consent records server-side with timestamp, IP, consent version, and method",
                "Provide a dedicated privacy settings page where users can withdraw individual consents",
                "Refresh consent if processing purpose changes significantly",
                "Avoid pre-checked boxes, confusing UX, or consent walls that block access",
                "Conduct a Legitimate Interest Assessment (LIA) as an alternative basis where consent is impractical",
            ],
            jurisdiction_notes="EU-wide requirement under GDPR. Individual member states may impose stricter requirements (e.g., Germany's TTDSG for cookies, France's CNIL guidelines).",
            confidence_level="high — based on GDPR text, CJEU case law, and supervisory authority guidance",
        )

    system = await _agent.build_system_prompt(request.user_id, request.organization_id)
    memory_context = request.metadata.get("memory_context", "")
    if memory_context:
        system += f"\n\n## Memory Context\n{memory_context}"
    raw = await _llm.complete(
        provider=_agent.default_provider, model=_agent.default_model,
        system=system,
        messages=[{"role": "user", "content": (
            f"Research this legal question:\n{request.query}\n\n"
            f"Jurisdiction: {request.jurisdiction}\n"
            f"Legal areas: {', '.join(request.legal_areas) if request.legal_areas else 'general'}\n\n"
            "Return ONLY a JSON object (no markdown fences) — be concise:\n"
            "summary (2-3 sentences), "
            "applicable_laws (list of strings, max 6), "
            "key_requirements (list of strings, max 6), "
            "relevant_cases (list of strings, max 4), "
            "practical_guidance (list of strings, max 5), "
            "jurisdiction_notes (1 sentence), "
            "confidence_level (exactly one word: high, medium, or low)"
        )}],
        max_tokens=1200,
    )
    tokens_used = _llm.count_tokens(raw)
    try:
        data = safe_json_loads(strip_json_fences(raw))
        # Guard: if summary looks like JSON (LLM nested the response), re-parse it
        if isinstance(data.get("summary"), str) and data["summary"].strip().startswith("{"):
            try:
                inner = safe_json_loads(data["summary"])
                if isinstance(inner, dict) and "summary" in inner:
                    data = inner
            except Exception:
                pass
        # Normalise confidence_level to a single word
        cl = str(data.get("confidence_level", "medium")).lower().split()[0]
        data["confidence_level"] = cl if cl in ("high", "medium", "low") else "medium"
        return LegalResearchResponse(**{k: v for k, v in data.items() if k in LegalResearchResponse.model_fields}, tokens_used=tokens_used, model_used=_agent.default_model)
    except Exception:
        return LegalResearchResponse(
            summary="Legal research completed. See guidance below.",
            applicable_laws=[], key_requirements=[], relevant_cases=[],
            practical_guidance=[raw[:400]] if raw else [],
            jurisdiction_notes=request.jurisdiction,
            confidence_level="medium",
            tokens_used=tokens_used,
            model_used=_agent.default_model,
        )


@router.post("/compliance-check", response_model=ComplianceCheckResponse, summary="Check regulatory compliance")
async def compliance_check(request: ComplianceCheckRequest) -> ComplianceCheckResponse:
    """Evaluate compliance against GDPR, CCPA, SOC2, HIPAA and other frameworks."""
    if settings.MOCK_MODE:
        return ComplianceCheckResponse(
            overall_status="partial",
            framework_results=[
                {
                    "framework": "GDPR",
                    "status": "non_compliant",
                    "gaps": [
                        "No explicit consent flow for data collection",
                        "EU data stored on US servers without valid transfer mechanism (SCCs or adequacy decision required)",
                        "90-day retention not documented or justified under storage limitation principle",
                    ],
                    "requirements": [
                        "Article 7 — valid consent mechanism",
                        "Chapter V — lawful EU-to-US data transfer mechanism",
                        "Article 5(1)(e) — storage limitation with defined retention policy",
                        "Article 13/14 — privacy notice with processing purposes",
                    ],
                },
                {
                    "framework": "CCPA",
                    "status": "partial",
                    "gaps": [
                        "No 'Do Not Sell or Share My Personal Information' opt-out link",
                        "Privacy notice does not describe categories of data collected or purposes",
                    ],
                    "requirements": [
                        "CCPA Section 1798.120 — right to opt out of sale/sharing",
                        "Section 1798.100 — right to know what data is collected",
                        "Section 1798.130 — designated methods for submitting requests",
                    ],
                },
            ],
            critical_gaps=[
                "Storing EU personal data on US servers without SCCs — active GDPR violation",
                "No consent mechanism — GDPR Article 7 violation with fines up to 4% of global revenue",
                "No opt-out mechanism for California residents — CCPA violation",
            ],
            remediation_steps=[
                {"priority": "high", "action": "Execute Standard Contractual Clauses (SCCs) for EU→US data transfers, or migrate EU data to EU-based infrastructure within 30 days"},
                {"priority": "high", "action": "Implement a consent management platform with GDPR-compliant granular consent flows for each processing purpose"},
                {"priority": "high", "action": "Add 'Do Not Sell or Share My Personal Information' link to footer and implement opt-out mechanism for California residents"},
                {"priority": "medium", "action": "Update privacy notice to describe data categories, purposes, retention periods, and user rights"},
                {"priority": "medium", "action": "Document and justify the 90-day retention period or reduce it to minimum necessary"},
            ],
            estimated_effort="2-4 weeks for critical compliance items, 2-3 months for full implementation and documentation",
        )

    system = await _agent.build_system_prompt(request.user_id, request.organization_id, use_brand_kit=False)
    memory_context = request.metadata.get("memory_context", "")
    if memory_context:
        system += f"\n\n## Memory Context\n{memory_context}"
    raw = await _llm.complete(
        provider=_agent.default_provider, model=_agent.default_model,
        system=system,
        messages=[{"role": "user", "content": (
            f"Evaluate the regulatory compliance of this practice or document:\n{request.description}\n\n"
            f"Business context: {request.business_context or 'Not provided'}\n"
            f"Check against: {', '.join(request.frameworks)}\n\n"
            "Return ONLY a JSON object (no markdown fences) with keys: "
            "overall_status (compliant/partial/non_compliant), "
            "framework_results (list of {framework, status, gaps, requirements}), "
            "critical_gaps (list of strings), "
            "remediation_steps (list of {priority: high/medium/low, action: string}), "
            "estimated_effort (string)"
        )}],
    )
    tokens_used = _llm.count_tokens(raw)
    try:
        data = json.loads(strip_json_fences(raw))
        return ComplianceCheckResponse(**data, tokens_used=tokens_used, model_used=_agent.default_model)
    except Exception:
        return ComplianceCheckResponse(
            overall_status="unknown",
            framework_results=[], critical_gaps=[],
            remediation_steps=[],
            estimated_effort="Manual review required",
            tokens_used=tokens_used,
            model_used=_agent.default_model,
        )


@router.post("/delete-source", response_model=DeleteSourceResponse, summary="Delete an ingested document")
async def delete_source(request: DeleteSourceRequest) -> DeleteSourceResponse:
    """Remove all RAG chunks for a (user_id, source_id) pair."""
    deleted = await _rag.delete_source(request.user_id, request.source_id)
    return DeleteSourceResponse(deleted_chunks=deleted)


@router.post("/sources", response_model=ListSourcesResponse, summary="List ingested documents")
async def list_sources(request: ListSourcesRequest) -> ListSourcesResponse:
    """List all ingested documents for a user under the lex agent."""
    rows = await _rag.list_sources(request.user_id, source_agent="lex")
    return ListSourcesResponse(
        sources=[SourceSummary(**r) for r in rows],
    )
