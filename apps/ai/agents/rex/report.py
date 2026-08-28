"""
Rex dataset → downloadable DOCX report.

Flow:
  1. profile_dataset        – pure Python, no LLM
  2. compute_kpis           – pure Python, deterministic
  3. plan_report_sections   – 1 LLM call → section list
  4. execute_section        – pandas aggregation + matplotlib chart PNG
  5. write_section_narrative – 1 LLM call per section (parallelized)
  6. write_takeaways        – 1 LLM call
  7. build_docx             – python-docx assembly
"""

import asyncio
import base64
import io
import json
from collections import Counter
from datetime import datetime

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ── Chart palette ──────────────────────────────────────────────────────────────

BRAND_GREEN = "#1DBC87"
BRAND_DARK  = "#0B2447"
BRAND_GRAY  = "#6B7280"

CHART_COLORS = [
    "#1DBC87", "#6366F1", "#f59e0b", "#ef4444",
    "#8b5cf6", "#06b6d4", "#ec4899", "#84cc16",
    "#f97316", "#14b8a6",
]


# ─── 1. Dataset profiling ─────────────────────────────────────────────────────

def profile_dataset(headers: list, rows: list, col_types: dict) -> dict:
    profile: dict = {"row_count": len(rows), "col_count": len(headers), "columns": []}
    for h in headers:
        ctype = col_types.get(h, "text")
        values = [r.get(h) for r in rows if r.get(h) not in (None, "", "N/A", "n/a")]
        info: dict = {"name": h, "type": ctype}
        if ctype == "numeric":
            nums: list[float] = []
            for v in values:
                try:
                    nums.append(float(
                        str(v).replace(",", "").replace("$", "").replace("₹", "")
                               .replace("%", "").strip()
                    ))
                except Exception:
                    pass
            if nums:
                info.update(
                    min=round(min(nums), 2), max=round(max(nums), 2),
                    mean=round(sum(nums) / len(nums), 2),
                    sum=round(sum(nums), 2), non_null=len(nums),
                )
        elif ctype in ("categorical", "text"):
            sv = [str(v) for v in values]
            cnt = Counter(sv)
            info["unique_count"] = len(cnt)
            info["top_values"] = [{"value": k, "count": c} for k, c in cnt.most_common(5)]
        elif ctype == "date":
            sv = sorted(str(v) for v in values if v)
            if sv:
                info["min_date"] = sv[0]
                info["max_date"] = sv[-1]
        profile["columns"].append(info)
    return profile


# ─── 2. KPI computation ───────────────────────────────────────────────────────

def compute_kpis(profile: dict) -> list[dict]:
    kpis: list[dict] = [
        {"label": "Total Rows", "value": f"{profile['row_count']:,}", "subtitle": "Records analyzed"}
    ]
    num_cols  = [c for c in profile["columns"] if c["type"] == "numeric" and "sum" in c]
    cat_cols  = [c for c in profile["columns"] if c["type"] in ("categorical", "text") and "unique_count" in c]
    date_cols = [c for c in profile["columns"] if c["type"] == "date"]

    if num_cols:
        c = num_cols[0]
        s = c["sum"]
        val = f"{s:,.0f}" if abs(s) >= 100 else f"{s:,.2f}"
        kpis.append({"label": f"Total {c['name']}", "value": val, "subtitle": f"Sum of {c['name']}"})
    if cat_cols:
        c = cat_cols[0]
        kpis.append({"label": f"Unique {c['name']}", "value": f"{c['unique_count']:,}",
                     "subtitle": f"Distinct {c['name']} values"})
    if num_cols:
        c = num_cols[0]
        kpis.append({"label": f"Avg {c['name']}", "value": f"{c['mean']:,.2f}",
                     "subtitle": f"Mean of {c['name']}"})
    elif date_cols:
        c = date_cols[0]
        kpis.append({"label": "Date Range",
                     "value": f"{c.get('min_date','?')} – {c.get('max_date','?')}",
                     "subtitle": f"{c['name']} coverage"})
    else:
        kpis.append({"label": "Columns", "value": str(profile["col_count"]), "subtitle": "Fields analyzed"})

    while len(kpis) < 4:
        kpis.append({"label": "Columns", "value": str(profile["col_count"]), "subtitle": "Fields analyzed"})
    return kpis[:4]


# ─── 3. LLM: plan sections ────────────────────────────────────────────────────

# Columns whose names suggest they are identifiers, free text, or contact info
# — not useful for aggregation charts.
_SKIP_KEYWORDS = {
    "email", "mail", "url", "website", "link", "address", "description",
    "desc", "note", "comment", "remarks", "keyword", "tag", "label",
    "phone", "mobile", "fax", "id", "uid", "uuid", "code", "ref",
    "no", "number", "num", "sr", "serial",
}

def _is_low_value_col(col_info: dict, row_count: int) -> bool:
    """Return True if a column is unlikely to produce a useful chart."""
    name_lower = col_info["name"].lower()
    # Skip if name contains identifier/contact/free-text keywords
    if any(kw in name_lower for kw in _SKIP_KEYWORDS):
        return True
    # Skip categorical columns where almost every value is unique (high cardinality)
    if col_info["type"] in ("categorical", "text"):
        unique = col_info.get("unique_count", 0)
        if row_count > 0 and unique / row_count > 0.5:
            return True
    return False


def _profile_lines(profile: dict) -> list[str]:
    lines: list[str] = []
    row_count = profile["row_count"]
    for c in profile["columns"]:
        if _is_low_value_col(c, row_count):
            continue  # don't expose low-value columns to the LLM planner
        if c["type"] == "numeric":
            lines.append(
                f"  - {c['name']} [numeric]: sum={c.get('sum','?')}, "
                f"mean={c.get('mean','?')}, min={c.get('min','?')}, max={c.get('max','?')}"
            )
        elif c["type"] in ("categorical", "text"):
            top = ", ".join(f"{t['value']}({t['count']})" for t in c.get("top_values", [])[:3])
            lines.append(f"  - {c['name']} [categorical]: {c.get('unique_count','?')} unique. Top: {top}")
        elif c["type"] == "date":
            lines.append(f"  - {c['name']} [date]: {c.get('min_date')} to {c.get('max_date')}")
    return lines


async def plan_report_sections(
    dataset_name: str,
    profiles: dict[str, dict],
    llm,
) -> list[dict]:
    is_multi = len(profiles) > 1

    sheet_blocks: list[str] = []
    all_valid: dict[str, set[str]] = {}
    for sheet_name, profile in profiles.items():
        lines = _profile_lines(profile)
        if not lines:
            continue
        header = f'Sheet "{sheet_name}" — {profile["row_count"]} rows:' if is_multi else f'Dataset — {profile["row_count"]} rows:'
        sheet_blocks.append(header + "\n" + "\n".join(lines))
        all_valid[sheet_name] = {c["name"] for c in profile["columns"]}

    if not sheet_blocks:
        return []

    columns_text = "\n\n".join(sheet_blocks)
    sheet_field = '  "sheet": "exact_sheet_name",\n' if is_multi else ""
    sheet_rule  = (
        f'- "sheet" must be one of: {list(profiles.keys())}\n'
        "- x_col and y_col must exist in the columns of that specific sheet\n"
        if is_multi else
        "- x_col and y_col must exist in the dataset columns above\n"
    )

    prompt = (
        f'Dataset: "{dataset_name}"\n\n'
        f"{columns_text}\n\n"
        'Plan 5-7 analytical sections for a professional business report. '
        'Return ONLY this JSON (no markdown):\n'
        '{"sections": [{\n'
        + sheet_field +
        '  "title": "Section title",\n'
        '  "subtitle": "One-line finding teaser",\n'
        '  "analysis_type": "top_n",\n'
        '  "x_col": "exact_column_name",\n'
        '  "y_col": "exact_numeric_column_or_null",\n'
        '  "agg_fn": "sum",\n'
        '  "chart_type": "horizontal_bar",\n'
        '  "top_n": 10,\n'
        '  "insight_hint": "what to highlight"\n'
        '}]}\n\n'
        "Rules:\n"
        "- analysis_type MUST be one of: top_n | distribution | time_series | concentration\n"
        "- chart_type MUST be one of: horizontal_bar | bar | pie | line\n"
        "- pie: ONLY when category has ≤ 6 meaningful distinct values. Set top_n ≤ 6. Otherwise use horizontal_bar\n"
        "- horizontal_bar: rankings and top-N; set top_n ≤ 10\n"
        "- SKIP columns that are IDs, serial numbers, emails, URLs, free-text descriptions, or keyword/tag fields\n"
        "- Focus on columns that reveal business patterns: who, what, where, how much\n"
        "- x_col must be a low-cardinality categorical (< 50 unique values) or date column\n"
        "- y_col must be numeric (or null for count-based)\n"
        + sheet_rule +
        "- Make sections diverse: top-N rankings, distributions, trends (if date col), geographic breakdowns\n"
        + ("- Spread sections across sheets so each sheet gets at least one section\n" if is_multi else "")
    )

    from core.utils import safe_json_loads
    raw = await llm.complete(
        provider="openai", model="gpt-5.6-luna",
        system="You plan professional business data reports. Return only valid JSON.",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        response_format={"type": "json_object"},
    )
    result = safe_json_loads(raw)
    sections = result.get("sections", []) if isinstance(result, dict) else []

    VALID_ANALYSIS = {"top_n", "distribution", "time_series", "concentration"}
    VALID_CHARTS   = {"bar", "horizontal_bar", "pie", "line"}
    default_sheet  = next(iter(profiles))
    row_count      = profiles[default_sheet]["row_count"]

    validated: list[dict] = []
    for s in sections:
        if not isinstance(s, dict) or not s.get("title"):
            continue
        sheet = s.get("sheet", default_sheet) if is_multi else default_sheet
        if sheet not in all_valid:
            sheet = default_sheet
        valid_cols    = all_valid[sheet]
        sheet_profile = profiles[sheet]
        x_col         = s.get("x_col", "")

        if x_col not in valid_cols:
            continue
        if s.get("analysis_type", "") not in VALID_ANALYSIS:
            continue
        if s.get("chart_type", "") not in VALID_CHARTS:
            continue
        if s.get("y_col") and s.get("y_col") not in valid_cols:
            continue

        # Reject low-value columns even if LLM tried to use them
        x_col_info = next((c for c in sheet_profile["columns"] if c["name"] == x_col), None)
        if x_col_info and _is_low_value_col(x_col_info, sheet_profile["row_count"]):
            continue

        s["sheet"] = sheet
        validated.append(s)
    return validated


def _fallback_sections(profiles: dict[str, dict]) -> list[dict]:
    sections: list[dict] = []
    for sheet_name, profile in profiles.items():
        num_cols  = [c for c in profile["columns"] if c["type"] == "numeric"]
        cat_cols  = [c for c in profile["columns"] if c["type"] in ("categorical", "text")
                     and not _is_low_value_col(c, profile["row_count"])]
        date_cols = [c for c in profile["columns"] if c["type"] == "date"]
        base = {"sheet": sheet_name}
        if cat_cols and num_cols:
            sections.append({**base,
                "title": f"Top 10 by {num_cols[0]['name']}", "subtitle": f"Highest {num_cols[0]['name']} by {cat_cols[0]['name']}",
                "analysis_type": "top_n", "x_col": cat_cols[0]["name"], "y_col": num_cols[0]["name"],
                "agg_fn": "sum", "chart_type": "horizontal_bar", "top_n": 10, "insight_hint": "top performers",
            })
        if date_cols and num_cols:
            sections.append({**base,
                "title": f"{num_cols[0]['name']} Over Time", "subtitle": "Trend over the dataset period",
                "analysis_type": "time_series", "x_col": date_cols[0]["name"], "y_col": num_cols[0]["name"],
                "agg_fn": "sum", "chart_type": "line", "top_n": None, "insight_hint": "trends and seasonality",
            })
        if not sections and profile["columns"]:
            col = next((c for c in profile["columns"] if not _is_low_value_col(c, profile["row_count"])),
                       profile["columns"][0])
            sections.append({**base,
                "title": f"Distribution of {col['name']}", "subtitle": "Value counts",
                "analysis_type": "distribution", "x_col": col["name"], "y_col": None,
                "agg_fn": "count", "chart_type": "bar", "top_n": 10, "insight_hint": "distribution",
            })
    return sections


# ─── 4. Execute section (pandas + matplotlib) ─────────────────────────────────

def _build_df(raw_table: dict) -> pd.DataFrame:
    headers   = raw_table.get("headers", [])
    rows      = raw_table.get("rows", [])
    col_types = raw_table.get("columnTypes", {})
    df = pd.DataFrame(rows, columns=headers) if rows else pd.DataFrame(columns=headers)
    for col, dtype in col_types.items():
        if col not in df.columns:
            continue
        if dtype == "numeric":
            df[col] = pd.to_numeric(
                df[col].astype(str)
                    .str.replace(",", "", regex=False).str.replace("$", "", regex=False)
                    .str.replace("₹", "", regex=False).str.replace("%", "", regex=False)
                    .str.strip(),
                errors="coerce",
            )
        elif dtype == "date":
            df[col] = pd.to_datetime(df[col], errors="coerce")
    return df


def execute_section(section: dict, dfs: dict[str, pd.DataFrame]) -> tuple[list[dict], bytes | None]:
    sheet_name    = section.get("sheet", "")
    df = dfs.get(sheet_name) if sheet_name and sheet_name in dfs else next(iter(dfs.values()), pd.DataFrame())

    x_col         = section.get("x_col", "")
    y_col         = section.get("y_col")
    agg_fn        = section.get("agg_fn", "sum")
    chart_type    = section.get("chart_type", "bar")
    analysis_type = section.get("analysis_type", "top_n")
    # Pie charts look bad with many slices — cap aggregation before rendering
    raw_top_n = section.get("top_n") or 15
    top_n = min(raw_top_n, 6) if chart_type == "pie" else raw_top_n

    title = section.get("title", "")

    if x_col not in df.columns:
        return [], None

    agg_data: list[dict] = []
    eff_y = y_col

    try:
        if analysis_type in ("top_n", "distribution", "concentration"):
            if y_col and y_col in df.columns:
                grp = df.groupby(x_col)[y_col].agg(agg_fn).reset_index()
                grp.columns = [x_col, y_col]
                grp = grp.dropna().sort_values(y_col, ascending=False)
            else:
                grp = df[x_col].value_counts().reset_index()
                grp.columns = [x_col, "count"]
                eff_y = "count"
            if analysis_type == "concentration" and eff_y:
                total = grp[eff_y].sum()
                if total:
                    grp["pct"] = (grp[eff_y] / total * 100).round(1)
            agg_data = grp.head(top_n).to_dict("records")

        elif analysis_type == "time_series":
            if y_col and y_col in df.columns:
                tmp = df[[x_col, y_col]].copy().dropna()
                tmp[x_col] = tmp[x_col].astype(str).str[:7]
                tmp = tmp.groupby(x_col)[y_col].sum().reset_index()
                agg_data = tmp.to_dict("records")

        else:
            if y_col and y_col in df.columns:
                grp = df.groupby(x_col)[y_col].agg("sum").reset_index()
                grp.columns = [x_col, y_col]
                grp = grp.dropna().sort_values(y_col, ascending=False)
                eff_y = y_col
            else:
                grp = df[x_col].value_counts().reset_index()
                grp.columns = [x_col, "count"]
                eff_y = "count"
            agg_data = grp.head(top_n).to_dict("records")

    except Exception:
        agg_data = []

    chart_png = _render_chart_png(chart_type, agg_data, x_col, eff_y or "count", title)
    return agg_data, chart_png


def _fmt_val(v: float) -> str:
    if abs(v) >= 1_000_000:
        return f"{v / 1_000_000:.1f}M"
    if abs(v) >= 1_000:
        return f"{v / 1_000:.1f}K"
    return f"{v:.1f}"


def _render_chart_png(chart_type: str, data: list[dict], x_key: str, y_key: str, title: str) -> bytes | None:
    if not data:
        return None
    try:
        fig, ax = plt.subplots(figsize=(8, 5.6))
        fig.patch.set_facecolor("white")
        ax.set_facecolor("white")

        x_vals = [str(row.get(x_key, ""))[:35] for row in data]
        y_vals = [float(row.get(y_key) or 0) for row in data]

        if chart_type == "horizontal_bar":
            colors = [CHART_COLORS[i % len(CHART_COLORS)] for i in range(len(x_vals))]
            bars = ax.barh(x_vals[::-1], y_vals[::-1], color=colors[::-1], alpha=0.90, height=0.58)
            ax.set_xlabel(y_key, fontsize=7, color=BRAND_GRAY)
            for bar, val in zip(bars, y_vals[::-1]):
                ax.text(bar.get_width() * 1.01 + 0.001, bar.get_y() + bar.get_height() / 2,
                        _fmt_val(val), va="center", fontsize=7, color=BRAND_DARK)
            ax.set_xlim(0, max(y_vals) * 1.20 if y_vals else 1)

        elif chart_type == "bar":
            colors = [CHART_COLORS[i % len(CHART_COLORS)] for i in range(len(x_vals))]
            bars = ax.bar(range(len(x_vals)), y_vals, color=colors, alpha=0.90, width=0.58)
            ax.set_xticks(range(len(x_vals)))
            ax.set_xticklabels(x_vals, rotation=35, ha="right", fontsize=7)
            for bar, val in zip(bars, y_vals):
                ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() * 1.02,
                        _fmt_val(val), ha="center", fontsize=6.5, color=BRAND_DARK)

        elif chart_type == "pie":
            # Cap at 6 slices, group tail as "Others"
            PIE_MAX = 6
            if len(x_vals) > PIE_MAX:
                keep_x = x_vals[:PIE_MAX - 1]
                keep_y = y_vals[:PIE_MAX - 1]
                x_vals = keep_x + ["Others"]
                y_vals = keep_y + [sum(y_vals[PIE_MAX - 1:])]
            wedges, texts, autotexts = ax.pie(
                y_vals, labels=x_vals, autopct="%1.1f%%",
                colors=CHART_COLORS[:len(y_vals)], startangle=90, pctdistance=0.80,
                wedgeprops={"linewidth": 0.5, "edgecolor": "white"},
            )
            for t in texts:     t.set_fontsize(7.5)
            for t in autotexts: t.set_fontsize(6.5)

        elif chart_type == "line":
            ax.plot(range(len(x_vals)), y_vals, color=BRAND_GREEN, linewidth=2.2,
                    marker="o", markersize=4.5, zorder=3)
            ax.fill_between(range(len(x_vals)), y_vals, alpha=0.08, color=BRAND_GREEN)
            step = max(1, len(x_vals) // 8)
            ax.set_xticks(range(0, len(x_vals), step))
            ax.set_xticklabels([x_vals[i] for i in range(0, len(x_vals), step)],
                               rotation=30, ha="right", fontsize=7)

        if chart_type != "pie":
            ax.spines["top"].set_visible(False)
            ax.spines["right"].set_visible(False)
            ax.spines["left"].set_color("#E5E7EB")
            ax.spines["bottom"].set_color("#E5E7EB")
            ax.tick_params(axis="both", colors=BRAND_GRAY, labelsize=7)
            ax.yaxis.grid(True, color="#F3F4F6", linewidth=0.6, zorder=0)
            ax.set_axisbelow(True)

        ax.set_title(title, fontsize=9, fontweight="bold", pad=8, color=BRAND_DARK)
        plt.tight_layout(pad=0.6)
        buf = io.BytesIO()
        fig.savefig(buf, format="png", dpi=140, bbox_inches="tight", facecolor="white")
        plt.close(fig)
        buf.seek(0)
        return buf.read()
    except Exception:
        plt.close("all")
        return None


# ─── 5. LLM: section narrative ────────────────────────────────────────────────

async def write_section_narrative(section: dict, agg_data: list[dict], llm) -> list[str]:
    if not agg_data:
        return []
    preview = json.dumps(agg_data[:12], default=str)
    prompt = (
        f'Chart: "{section["title"]}"\n'
        f"Data:\n{preview}\n\n"
        f'Write 3-5 concise bullet insights using exact numbers from the data. '
        f'Focus on: {section.get("insight_hint", "key patterns")}.\n'
        'Return JSON: {"bullets": ["insight 1", "insight 2", ...]}'
    )
    from core.utils import safe_json_loads
    raw = await llm.complete(
        provider="openai", model="gpt-5.6-luna",
        system="You are a data analyst. Return only valid JSON.",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        response_format={"type": "json_object"},
    )
    result = safe_json_loads(raw)
    if isinstance(result, dict):
        bullets = result.get("bullets", [])
        if isinstance(bullets, list):
            return [b for b in bullets[:5] if b]
    return []


# ─── 6. LLM: key takeaways ────────────────────────────────────────────────────

async def write_takeaways(
    dataset_name: str,
    sections: list[dict],
    all_narratives: list[list[str]],
    llm,
) -> list[dict]:
    summary = "\n".join(
        f"- {s['title']}: " + "; ".join(n[:2])
        for s, n in zip(sections, all_narratives)
        if n
    )
    prompt = (
        f'Dataset: "{dataset_name}"\n\nFindings:\n{summary}\n\n'
        "Write exactly 5 strategic takeaways. "
        'Return JSON: {"takeaways": [{"title": "Short bold label", "body": "2-sentence insight with specific numbers"}]}'
    )
    from core.utils import safe_json_loads
    raw = await llm.complete(
        provider="openai", model="gpt-5.6-luna",
        system="You are a strategic analyst. Return only valid JSON.",
        messages=[{"role": "user", "content": prompt}],
        temperature=0.3,
        response_format={"type": "json_object"},
    )
    result = safe_json_loads(raw)
    if isinstance(result, dict):
        tws = result.get("takeaways", [])
        if isinstance(tws, list):
            return tws[:5]
    return []


# ─── 7. DOCX assembly ─────────────────────────────────────────────────────────

def _add_heading(doc: Document, text: str, level: int, color_hex: str = "0B2447") -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        )


def build_docx(
    dataset_name: str,
    org_name: str,
    kpis: list[dict],
    sections: list[dict],
    section_charts: list[bytes | None],
    section_narratives: list[list[str]],
    takeaways: list[dict],
    generated_at: str,
) -> bytes:
    doc = Document()

    # ── Cover ────────────────────────────────────────────────────────────────
    h = doc.add_heading(dataset_name, 0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sub = doc.add_paragraph("Data Analysis Report")
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0x1D, 0xBC, 0x87)
    if org_name:
        doc.add_paragraph(f"Prepared for: {org_name}").runs[0].font.size = Pt(10)
    doc.add_paragraph(f"Generated: {generated_at}").runs[0].font.size = Pt(10)
    doc.add_page_break()

    # ── Executive Summary ────────────────────────────────────────────────────
    _add_heading(doc, "Executive Summary", 1)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    for i, kpi in enumerate(kpis[:4]):
        cell = tbl.rows[0].cells[i]
        p = cell.paragraphs[0]
        run = p.add_run(kpi["value"])
        run.bold = True
        run.font.size = Pt(16)
        label = cell.add_paragraph(kpi["label"])
        label.runs[0].font.size = Pt(9)
        label.runs[0].bold = True
        sub_p = cell.add_paragraph(kpi["subtitle"])
        sub_p.runs[0].font.size = Pt(8)
    doc.add_page_break()

    # ── Section pages ────────────────────────────────────────────────────────
    for sec, chart_png, bullets in zip(sections, section_charts, section_narratives):
        _add_heading(doc, sec.get("title", ""), 1)
        sub_p = doc.add_paragraph(sec.get("subtitle", ""))
        sub_p.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        sub_p.runs[0].font.size = Pt(10)

        if chart_png:
            doc.add_picture(io.BytesIO(chart_png), width=Inches(6.0))

        for b in (bullets or []):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(b).font.size = Pt(10)

        doc.add_page_break()

    # ── Key Takeaways ────────────────────────────────────────────────────────
    _add_heading(doc, "Key Takeaways", 1)
    doc.add_paragraph("Strategic recommendations based on the analysis").runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    for i, tw in enumerate(takeaways[:5]):
        num_p = doc.add_paragraph()
        run = num_p.add_run(f"{i + 1:02d}.  {tw.get('title', '')}")
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1D, 0xBC, 0x87)
        body_p = doc.add_paragraph(tw.get("body", ""))
        body_p.runs[0].font.size = Pt(10)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── 8. Main orchestrator ─────────────────────────────────────────────────────

async def generate_full_report(
    dataset_name: str,
    raw_table: dict,
    fmt: str,
    org_id: str,
    llm,
) -> tuple[str, str]:
    generated_at = datetime.utcnow().strftime("%B %d, %Y")

    # Build per-sheet raw tables
    raw_sheets: dict[str, dict] = {}
    if raw_table.get("sheets"):
        for sname, sdata in raw_table["sheets"].items():
            raw_sheets[sname] = sdata
    else:
        raw_sheets["Dataset"] = raw_table

    # Profile each sheet
    profiles: dict[str, dict] = {
        name: profile_dataset(
            s.get("headers", []), s.get("rows", []), s.get("columnTypes", {})
        )
        for name, s in raw_sheets.items()
    }

    # KPIs from primary sheet
    primary_profile = next(iter(profiles.values()))
    kpis = compute_kpis(primary_profile)

    # Plan sections (1 LLM call)
    sections = await plan_report_sections(dataset_name, profiles, llm)
    if not sections:
        sections = _fallback_sections(profiles)

    # Execute sections (CPU-bound → executor)
    loop = asyncio.get_event_loop()

    def _sync_execute() -> list[tuple[list[dict], bytes | None]]:
        dfs = {name: _build_df(s) for name, s in raw_sheets.items()}
        return [execute_section(sec, dfs) for sec in sections]

    section_results  = await loop.run_in_executor(None, _sync_execute)
    section_agg_data = [r[0] for r in section_results]
    section_charts   = [r[1] for r in section_results]

    # Drop sections with no data/chart
    valid = [
        (sec, agg, chart)
        for sec, agg, chart in zip(sections, section_agg_data, section_charts)
        if agg and chart
    ]
    if valid:
        sections, section_agg_data, section_charts = map(list, zip(*valid))
    else:
        sections, section_agg_data, section_charts = _fallback_sections(profiles), [], []

    # Narratives (parallelized)
    all_narratives: list[list[str]] = list(await asyncio.gather(
        *[write_section_narrative(sec, agg, llm)
          for sec, agg in zip(sections, section_agg_data)]
    ))

    # Takeaways (1 LLM call)
    takeaways = await write_takeaways(dataset_name, sections, all_narratives, llm)

    # Org name from brand kit
    org_name = ""
    try:
        from core.brand_kit import load_brand_kit
        bk = await load_brand_kit(org_id)
        org_name = bk.company_name or ""
    except Exception:
        pass

    # Build DOCX (CPU-bound → executor)
    def _sync_build() -> bytes:
        return build_docx(dataset_name, org_name, kpis, sections, section_charts,
                          all_narratives, takeaways, generated_at)

    raw_bytes = await loop.run_in_executor(None, _sync_build)

    mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    return base64.b64encode(raw_bytes).decode(), mime
